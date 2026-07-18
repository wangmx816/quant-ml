#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成 TASK5 Word/PDF：理论 + 五只股票分别建模实证。"""

from __future__ import annotations

import argparse
import json
import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(ROOT))

from src.train import run_all

CHART_DIR = ROOT / "output" / "charts"
METRICS_PATH = ROOT / "output" / "metrics.json"
TASK5_DIR = ROOT.parent / "TASK5"
REPO_URL = "https://github.com/wangmx816/quant-ml"
PAGES_URL = "https://wangmx816.github.io/quant-ml/"


def set_run_font(run, size_pt: float = 10.5, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def add_paragraph(doc: Document, text: str, first_line_indent: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, bold=True)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run)


def add_figure(doc: Document, path: Path, caption: str, interpretation: str) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Cm(14.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption)
    add_paragraph(doc, interpretation)


def build_doc(meta: dict, name: str) -> Path:
    stocks = meta["stocks"]
    gw = next(s for s in stocks if s["symbol"] == "002202")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    add_heading(doc, "机器学习分类算法与模型评价实验报告（TASK5）")
    add_paragraph(doc, f"姓名：{name}", first_line_indent=False)

    add_heading(doc, "一、分类型机器学习算法概述")
    add_paragraph(
        doc,
        "分类是监督学习的核心任务之一，目标是根据输入特征将样本映射到离散类别。"
        "本实验对金风科技、三一重工、徐工机械、安彩高科、智慧农业五只股票分别建立二分类模型，"
        "预测下一交易日收益是否为正，并比较逻辑回归、决策树与随机森林的表现。",
    )
    add_heading(doc, "1.1 逻辑回归")
    add_paragraph(
        doc,
        "逻辑回归对特征做线性加权后经 Sigmoid 映射为 (0,1) 概率，再按阈值判别类别。"
        "优点是可解释、训练快、概率输出稳定；缺点是难以刻画复杂非线性边界，实践中常配合标准化。",
    )
    add_heading(doc, "1.2 决策树")
    add_paragraph(
        doc,
        "决策树递归选择最优分裂特征，将样本划分到叶节点并输出多数类。"
        "优点是规则直观、能捕捉非线性与交互；缺点是单棵树易过拟合，需限制深度或剪枝。",
    )
    add_heading(doc, "1.3 随机森林")
    add_paragraph(
        doc,
        "随机森林通过 Bootstrap 样本与随机特征子集训练多棵树，再投票集成。"
        "优点是通常比单树更稳健、可输出特征重要性；缺点是可解释性较弱、模型体积更大。",
    )

    add_heading(doc, "二、机器学习模型评价指标")
    add_heading(doc, "2.1 混淆矩阵")
    add_paragraph(
        doc,
        "混淆矩阵统计真实标签与预测标签的对应关系，包含 TP、FP、TN、FN 四个单元，"
        "由此可计算准确率、精确率、召回率与 F1，能揭示漏报与误报的权衡。",
    )
    add_heading(doc, "2.2 ROC 曲线")
    add_paragraph(
        doc,
        "ROC 以假正率 FPR 为横轴、真正率 TPR 为纵轴，描述不同阈值下的分类表现。"
        "曲线越靠近左上角，模型区分能力越强；对角线对应随机猜测。",
    )
    add_heading(doc, "2.3 AUC")
    add_paragraph(
        doc,
        "AUC 为 ROC 曲线下面积，取值约在 0.5～1。越接近 1 表示正负样本排序能力越强，"
        "且对单一阈值不敏感，适合作为本实验的核心比较指标。",
    )

    add_heading(doc, "三、数据准备与实验设计")
    add_paragraph(
        doc,
        "参考 model_data.csv 的分类建模思路，并与 quant-strategy 项目保持相同的五只标的。"
        f"因原 model_data.csv 最新仅至约 2022 年，本实验改为拉取 "
        f"{meta['date_range']} 的前复权日线数据（东方财富），"
        "构造动量、波动、均线比、量比、RSI 等特征；"
        "应变量为下一交易日收益是否为正（Next_Ret>0→1）。"
        "重要原则：五只股票各自独立划分训练集/测试集并分别训练，不把截面样本混在一起。",
    )
    lines = []
    for s in stocks:
        lines.append(
            f"{s['name']}（{s['symbol']}）样本 {s['n_samples']} 条"
            f"（{s['date_min']}～{s['date_max']}），正类占比 {s['pos_rate']:.1%}，"
            f"最优模型 {s['best_model']}（AUC={s['best_auc']:.3f}）"
        )
    add_paragraph(doc, "各股票概况：" + "；".join(lines) + "。")
    add_paragraph(
        doc,
        f"划分与验证：{meta['split']}。"
        "逻辑回归配合标准化；决策树限制深度；随机森林 200 棵树。"
        f"交互看板与代码已发布：{REPO_URL} ，在线页 {PAGES_URL} 。",
    )

    add_heading(doc, "四、分股票建模结果")
    add_figure(
        doc,
        CHART_DIR / "fig1_auc_heatmap.png",
        "图1  五只股票 × 三类模型 测试集 AUC",
        "图1展示每只股票在三类模型上的测试集 AUC。"
        "可见不同标的的可预测性存在差异，同一模型在不同股票上的表现也不相同，"
        "这正说明“分别建模”比混合截面更符合单标的交易场景。",
    )
    add_figure(
        doc,
        CHART_DIR / "fig2_best_auc.png",
        "图2  各股票最优分类模型 AUC",
        "图2汇总每只股票的最优模型及其 AUC。"
        "灰色虚线为随机基准 0.5；高于该线越多，说明相对随机猜测的提升越明显。",
    )
    add_figure(
        doc,
        CHART_DIR / f"roc_{gw['symbol']}.png",
        "图3  金风科技三类模型 ROC 曲线",
        f"图3以金风科技为例展示 ROC。"
        f"其最优模型为 {gw['best_model']}（AUC={gw['best_auc']:.3f}）。"
        "其余四只股票的 ROC 图见输出目录 output/charts/roc_*.png，看板中可切换查看。",
    )

    add_heading(doc, "五、结论")
    add_paragraph(
        doc,
        "本报告完成了分类算法与评价指标的理论梳理，并在 2024-03-31 至 2026-06-30 区间内，"
        "对五只股票分别训练逻辑回归、决策树与随机森林，计算 AUC 并绘制 ROC。"
        "实证表明：分股票建模能够反映标的异质性；日频特征下样本量显著高于旧版季度面板，"
        "评价结果更稳定。后续可引入估值日频指标、宏观因子或时间序列交叉验证以进一步提升稳健性。",
    )

    out_docx = TASK5_DIR / f"{name}TASK5.docx"
    TASK5_DIR.mkdir(parents=True, exist_ok=True)
    # 同步一份到仓库根目录
    local_docx = ROOT / f"{name}TASK5.docx"
    doc.save(str(local_docx))
    doc.save(str(out_docx))
    return out_docx


def docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        d = word.Documents.Open(str(docx_path.resolve()))
        d.SaveAs(str(pdf_path.resolve()), FileFormat=17)
        d.Close(False)
    finally:
        word.Quit()


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--name", default=os.environ.get("STUDENT_NAME", "wangmx"))
    parser.add_argument("--skip-train", action="store_true")
    args = parser.parse_args()

    if args.skip_train and METRICS_PATH.exists():
        with open(METRICS_PATH, encoding="utf-8") as f:
            meta = json.load(f)
    else:
        meta = run_all()

    docx_path = build_doc(meta, args.name)
    pdf_path = TASK5_DIR / f"{args.name}TASK5.pdf"
    pdf_local = ROOT / f"{args.name}TASK5.pdf"
    try:
        docx_to_pdf(docx_path, pdf_path)
        # 复制到仓库
        import shutil

        shutil.copy2(pdf_path, pdf_local)
        print(f"PDF: {pdf_path}")
    except Exception as e:
        print(f"PDF 转换失败: {e}")
        print(f"Word: {docx_path}")


if __name__ == "__main__":
    main()
