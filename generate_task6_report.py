#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成 TASK6 Word 报告：含净值差距解读、等权差距解读、基线/改进对比。"""

from __future__ import annotations

import argparse
import json
import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(ROOT))

from src.strategy import run_all_strategies

CHART_DIR = ROOT / "output" / "charts"
STRATEGY_PATH = ROOT / "output" / "strategy_metrics.json"
TASK6_DIR = ROOT.parent / "TASK6"
REPO_URL = "https://github.com/wangmx816/quant-ml"
PAGES_URL = "https://wangmx816.github.io/quant-ml/"


def set_run_font(run, size_pt: float = 10.5, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def add_paragraph(doc: Document, text: str, first_line_indent: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, bold=True)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run)


def add_figure(doc: Document, path: Path, caption: str, interpretation: str) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Cm(14.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption)
    add_paragraph(doc, interpretation)


def build_doc(meta: dict, name: str) -> Path:
    baseline = meta["baseline"]
    improved = meta["improved"]
    analysis = meta.get("analysis", {})
    g = analysis.get("goldwind_gap", {})
    e = analysis.get("equal_weight_gap", {})

    gw_b = next(s for s in baseline["stocks"] if s["symbol"] == "002202")
    gw_i = next(s for s in improved["stocks"] if s["symbol"] == "002202")
    best_b, best_i = gw_b["best_model"], gw_i["best_model"]
    mb = gw_b["models"][best_b]["metrics"]
    mi = gw_i["models"][best_i]["metrics"]

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    add_heading(doc, "基于机器学习的量化交易策略实验报告（TASK6）")
    add_paragraph(doc, f"姓名：{name}", first_line_indent=False)
    add_paragraph(
        doc,
        f"代码与看板：{REPO_URL} ；在线看板：{PAGES_URL} "
        "（含“分类评价 / 交易策略”双 Tab，策略 Tab 可切换基线版与改进版）。",
        first_line_indent=False,
    )

    add_heading(doc, "一、机器学习交易策略的核心理念与优缺点")
    add_paragraph(
        doc,
        "基于机器学习的交易策略，核心理念是：用历史可观测特征（因子）学习价格未来方向或收益的统计规律，"
        "再把模型输出转化为可执行仓位规则。它强调数据驱动与可回测。"
        "优点是可融合多维因子、刻画非线性；缺点是噪声大、易过拟合，且必须严格按时间划分避免前视偏差。",
    )

    add_heading(doc, "二、常见自变量因子与应变量定义")
    add_paragraph(
        doc,
        "自变量包括动量（1/5/10/20日收益）、波动率、均线比、量比、RSI，以及改进版中的中期趋势特征"
        "（10/60均线比、价格相对60日均线、上涨日占比等），并视情况加入 PB、PS、市值。"
        "应变量：基线用“次日收益>0”；改进版用“未来5日收益>0”，更贴近趋势持仓。",
    )

    add_heading(doc, "三、实验设计：基线版与改进版")
    add_paragraph(
        doc,
        "数据为五只股票（与 quant-strategy 一致）2024-03-31～2026-06-30 日前复权日线，按时间前70%训练、后30%测试。"
        "基线版：开仓阈值0.55、预测次日涨跌、最低持仓1日。"
        "改进版：在0.50阈值思路基础上，因部分模型绝对概率整体偏低，改为按概率分位数开仓"
        "（默认做多约一半交易日，均线趋势向上时提升至约六成），"
        "标签改为未来5日收益方向，加入趋势特征，信号触发后最低持仓5日。"
        "手续费万三、滑点万一。看板可切换版本对比。",
    )

    add_heading(doc, "四、金风科技：策略净值为何远低于买入持有（基线解读）")
    add_paragraph(
        doc,
        f"测试区间为 {g.get('test_period', gw_b['test_start']+'~'+gw_b['test_end'])}。"
        f"基线最优模型为【{best_b}】：策略累计 {g.get('baseline_cum', mb['cumulative_return'])}%，"
        f"买入持有 {g.get('baseline_bench', mb['benchmark_cumulative'])}%，"
        f"超额 {g.get('baseline_excess', mb['excess_return'])}%，"
        f"持仓天数占比仅约 {100 * float(g.get('baseline_long_ratio', mb.get('long_ratio', 0))):.1f}%。",
    )
    add_paragraph(
        doc,
        "差距的本质不是“基准算错”，而是：测试期金风科技出现显著上涨行情，买入持有充分享受趋势；"
        "而基线策略因阈值偏严、标签过短，给出的开仓信号很少，资金大部分时间以现金形式存在。"
        "因此策略净值曲线更平滑、回撤更小，但绝对收益大幅落后——这是典型的“择时空仓踏空”。"
        "夏普较高只说明单位波动下的收益尚可，并不等于战胜买入持有。",
    )
    add_figure(
        doc,
        CHART_DIR / "task6_fig1_equity.png",
        "图1  金风科技基线版：三类模型净值 vs 买入持有",
        "图1可见买入持有终点净值明显高于各策略曲线。策略之间差异小于“策略相对基准”的差异，"
        "说明主要矛盾是仓位过低，而非单一模型选错。",
    )
    add_figure(
        doc,
        CHART_DIR / "task6_fig2_quarterly.png",
        "图2  金风科技基线最优模型分季度收益",
        "图2显示部分季度策略收益接近0（空仓），而基准在上涨季度录得较高正收益，进一步印证踏空。",
    )

    add_heading(doc, "五、等权组合与等权买入持有的差距是什么意思")
    add_paragraph(
        doc,
        "“等权组合”：对五只股票各自用其最优模型得到测试集净值曲线，再按日期对齐后取算术平均，相当于五策略各占20%资金。"
        "“等权买入持有”：同样五只股票，各自全程持有的净值曲线等权平均，相当于五只股票各买20%并一直拿着。"
        "二者的差距（超额）= 组合累计收益 − 等权买入持有累计收益。",
    )
    add_paragraph(
        doc,
        f"基线结果：等权组合累计约 {e.get('baseline_port_cum')}%，等权买入持有约 {e.get('baseline_bench_cum')}%，"
        f"超额约 {e.get('baseline_excess')}%；成分股平均持仓占比约 "
        f"{100 * float(e.get('baseline_avg_long') or 0):.1f}%。"
        "若超额为正，通常表示：策略在部分个股上通过空仓规避了下跌，或整体回撤更小，使得相对等权持有更有优势；"
        "但也可能伴随“上涨市踏空”，导致绝对收益并不高。"
        "基线等权组合正是这类情形——相对基准超额尚可，但绝对收益接近持平，因为多只股票也长期低仓位。",
    )
    add_figure(
        doc,
        CHART_DIR / "task6_fig5_bonus.png",
        "图5  基线版五股等权组合 vs 等权买入持有",
        "图5中若紫色组合线终点高于灰色基准，表示相对等权持有有超额；若两者都接近1，则说明策略整体仓位偏低、绝对增值有限。",
    )

    add_heading(doc, "六、参数改进：降阈值、拉长持仓、趋势标签/特征")
    add_paragraph(
        doc,
        "针对踏空问题，改进版做了三类调整：（1）以0.50为阈值参考，并改用概率分位数开仓"
        "（约50%–60%交易日持仓），解决绝对概率过低导致几乎空仓的问题；"
        "（2）标签由次日改为未来5日收益方向，并设置最低持仓5日；"
        "（3）增加20日动量、10/60均线比、价格/60日均线、趋势持续度等偏趋势特征。"
        f"改进后金风最优模型为【{best_i}】：累计 {g.get('improved_cum', mi['cumulative_return'])}%，"
        f"买入持有 {g.get('improved_bench', mi['benchmark_cumulative'])}%，"
        f"超额 {g.get('improved_excess', mi['excess_return'])}%，"
        f"持仓占比约 {100 * float(g.get('improved_long_ratio', mi.get('long_ratio', 0))):.1f}%"
        f"（基线仅约 {100 * float(g.get('baseline_long_ratio', 0)):.1f}%）。",
    )
    add_paragraph(
        doc,
        f"等权组合改进后：累计约 {e.get('improved_port_cum')}%，等权买入持有约 {e.get('improved_bench_cum')}%，"
        f"超额约 {e.get('improved_excess')}%，平均持仓占比约 "
        f"{100 * float(e.get('improved_avg_long') or 0):.1f}%。"
        "改进目标是缩小与买入持有的绝对收益差距，同时保留一定择时能力；是否全面优于基线，需结合夏普与回撤综合判断。",
    )
    add_figure(
        doc,
        CHART_DIR / "task6_fig6_version_compare.png",
        "图6  金风科技：基线策略 vs 改进策略 vs 买入持有",
        "图6直接对比两版最优策略净值。改进版持仓更积极时，曲线通常更接近买入持有；"
        "若仍明显低于基准，说明趋势市中完全战胜买入持有仍困难，但踏空程度应有所缓解。",
    )
    add_figure(
        doc,
        CHART_DIR / "task6_fig7_bonus_compare.png",
        "图7  五股等权：基线 vs 改进 vs 等权买入持有",
        "图7从组合层面观察参数改进效果。改进版若抬升紫色净值并改变相对灰色基准的位置，"
        "说明降阈值与拉长持仓改变了组合风险暴露。",
    )
    add_figure(
        doc,
        CHART_DIR / "task6_fig8_improved_equity.png",
        "图8  金风科技改进版三类模型净值对比",
        "图8给出改进参数下逻辑回归、决策树、随机森林的测试集净值，便于比较模型在趋势设定下的相对表现。",
    )

    add_heading(doc, "七、决策树与随机森林等效果对比")
    lines_b = []
    for mname, mres in gw_b["models"].items():
        x = mres["metrics"]
        lines_b.append(
            f"{mname}（基线）：累计{x['cumulative_return']}%/夏普{x['sharpe_ratio']}/持仓占比{100*x.get('long_ratio',0):.1f}%"
        )
    lines_i = []
    for mname, mres in gw_i["models"].items():
        x = mres["metrics"]
        lines_i.append(
            f"{mname}（改进）：累计{x['cumulative_return']}%/夏普{x['sharpe_ratio']}/持仓占比{100*x.get('long_ratio',0):.1f}%"
        )
    add_paragraph(doc, "；".join(lines_b) + "。")
    add_paragraph(doc, "；".join(lines_i) + "。")
    add_paragraph(
        doc,
        "评价时需同时看绝对收益、超额收益、回撤与夏普：高夏普低仓位可能“稳健但踏空”；"
        "高仓位更接近买入持有，则回撤往往上升。改进版意图在二者之间取得更合理的折中。",
    )
    add_figure(
        doc,
        CHART_DIR / "task6_fig3_compare.png",
        "图3–图4  基线版模型指标与五股最优夏普",
        "左图为金风基线三模型夏普与年化；右图为五股各自最优夏普，反映标的可交易性差异。",
    )

    add_heading(doc, "八、结论")
    add_paragraph(
        doc,
        "本报告说明了机器学习交易策略理念、因子与标签定义，并完成基线与改进两套回测。"
        "金风科技策略净值显著低于买入持有，主因是基线阈值偏严导致长期空仓踏空；"
        "等权组合相对等权买入持有的差距，衡量的是“五策略分散择时”相对“五股一直持有”的超额，"
        "正超额并不自动等于高绝对收益。"
        "通过降低阈值、5日趋势标签、趋势特征与最低持仓5日，改进版提高了持仓参与度并改变了净值路径。"
        f"完整交互结果见看板 {PAGES_URL} ，仓库 {REPO_URL} 。",
    )

    TASK6_DIR.mkdir(parents=True, exist_ok=True)
    out_docx = TASK6_DIR / f"{name}TASK6.docx"
    local_docx = ROOT / f"{name}TASK6.docx"
    doc.save(str(out_docx))
    doc.save(str(local_docx))
    return out_docx


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--name", default=os.environ.get("STUDENT_NAME", "王茗萱"))
    parser.add_argument("--skip-run", action="store_true")
    args = parser.parse_args()

    if args.skip_run and STRATEGY_PATH.exists():
        meta = json.loads(STRATEGY_PATH.read_text(encoding="utf-8"))
        if "baseline" not in meta:
            meta = run_all_strategies()
    else:
        meta = run_all_strategies()

    docx_path = build_doc(meta, args.name)
    print(f"Word: {docx_path}")


if __name__ == "__main__":
    main()
